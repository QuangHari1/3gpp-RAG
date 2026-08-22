"""Export the paper vocabulary's term definitions into a portable JSON asset."""

from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_SOURCE = PROJECT_ROOT / "newbaseline/resources/3GPP_vocabulary.docx"
DEFAULT_OUTPUT = PROJECT_ROOT / "newbaseline/resources/3GPP_definitions_paper_v17.json"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for block in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def extract_terms(path: Path) -> tuple[str, list[dict[str, str]]]:
    """Match the paper parser's term section and duplicate-last semantics."""
    document = Document(path)
    title = next((paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()), "")
    terms: dict[str, str] = {}
    references = 0
    in_terms = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if "References" in text:
            references += 1
        if references < 2:
            continue
        if "Terms and definitions" in text:
            in_terms = True
        elif "Abbreviations" in text:
            in_terms = False
        elif in_terms and ":" in text:
            term, definition = text.split(":", 1)
            terms[term.strip()] = definition.strip().rstrip(".")
    return title, [{"term": term, "definition": definition} for term, definition in sorted(terms.items())]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    title, terms = extract_terms(args.source)
    payload = {
        "schema_version": 1,
        "source_document": title,
        "source_file": args.source.name,
        "source_sha256": sha256(args.source),
        "term_count": len(terms),
        "terms": terms,
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"Exported {len(terms)} definitions to {args.output}")


if __name__ == "__main__":
    main()
